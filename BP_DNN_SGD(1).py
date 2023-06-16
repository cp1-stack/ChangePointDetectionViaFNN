# -*- coding: utf-8 -*-
import numpy as np
import argparse
from sklearn.metrics import accuracy_score
import numpy.random as r
import matplotlib.pyplot as plt
import struct
import random as rd
import math
from threading import Thread
import os
import pandas as pd
import docx
from docx import Document
from docx.shared import Cm
from time import *





X_trains_norm=data_normalize(np.transpose(X_trains))
X_tests_norm=data_normalize(np.transpose(X_tests))

def convert_y_to_vect(y):
    y_vect = np.zeros((len(y), 10))
    for i in range(len(y)):
        y_vect[i, np.int16(y[i])] = 1
    return y_vect
#y_v_train = convert_y_to_vect(train_labels)
#y_v_test = convert_y_to_vect(test_labels)

y_v_train = np.transpose(convert_y_to_vect(train_labels))
y_v_test = np.transpose(convert_y_to_vect(test_labels))
nn_structure = [784,1024,120,10]

def f(x):
    return 1.0 / (1 + np.exp(-x))

'''

def sigmoid(in_x):  #  RuntimeWarning: overflow encountered in exp
    # return 1.0/(1 + np.exp(-in_x))
    # 优化方法
    if in_x >= 0:
        return 1.0/(1+np.exp(-in_x))
    else:
        return np.exp(in_x)/(1+np.exp(in_x))
'''
def f_deriv(x):
    return f(x) * (1 - f(x))
def ReluFunc(x):
    x=(np.abs(x)+x)/2.0
    return x
def ReluPrime(x):
    x[x<=0]=0
    x[x>0]=1
    return x
def LReluFunc(x):
    x[x<=0]=0.01*x
    x[x>0]=x
    return x
def LReluPrime(x):
    x[x<=0]=0.01
    x[x>0]=1
    return x

def random_mini_batches(X, Y, mini_batch_size=64):
    m = X.shape[1]  # number of training examples
    mini_batches = []
    # Step 1: Shuffle (X, Y)
    permutation = list(np.random.permutation(m))#np.random.permutation()随机排列序列
    shuffled_X = X[:, permutation]
    shuffled_Y = Y[:, permutation]
    # Step 2: Partition (shuffled_X, shuffled_Y). Minus the end case.
    num_complete_minibatches = math.floor(m / mini_batch_size)  # number of mini batches of size mini_batch_size in your partitionning
    for k in range(0, num_complete_minibatches):
        mini_batch_X = shuffled_X[:, k * mini_batch_size: (k + 1) * mini_batch_size]
        mini_batch_Y = shuffled_Y[:, k * mini_batch_size: (k + 1) * mini_batch_size]
        mini_batch = (mini_batch_X, mini_batch_Y)
        mini_batches.append(mini_batch)
    if m % mini_batch_size != 0:
        mini_batch_X = shuffled_X[:, (k + 1) * mini_batch_size:]
        mini_batch_Y = shuffled_Y[:, (k + 1) * mini_batch_size:]
        mini_batch = (mini_batch_X, mini_batch_Y)
        mini_batches.append(mini_batch)
    return mini_batches


def setup_and_init_weights(nn_structure):
    W = {}
    b = {}
    #np.random.seed(3)
    for l in range(1, len(nn_structure)):
        #W[l] = np.random.randn(nn_structure[l], nn_structure[l-1])*np.sqrt(1/nn_structure[l-1])
        #b[l] = np.zeros((nn_structure[l],))
        W[l] = np.random.randn(nn_structure[l], nn_structure[l - 1]) * np.sqrt(1 / nn_structure[l - 1])
        b[l] = np.zeros((nn_structure[l],1))
    return W, b

def init_tri_values(nn_structure):
    tri_W = {}
    tri_b = {}
    for l in range(1, len(nn_structure)):
        tri_W[l] = np.zeros((nn_structure[l], nn_structure[l-1]))
        tri_b[l] = np.zeros((nn_structure[l],1))
    return tri_W, tri_b

def update_parameters_with_bsgd(W, b, tri_W, tri_b, alpha, lamb):
    # Update rule for each parameter
    for l in range(len(nn_structure) - 1, 0, -1):
        W[l] += -alpha * (tri_W[l] + lamb * W[l])
        b[l] += -alpha * tri_b[l]
    return W,b

def initialize_with_momentum(nn_structure):
     # number of layers in the neural networks
    v = {}
    for l in range(1,len(nn_structure)):
        v["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        v["db" + str(l)] = np.zeros((nn_structure[l],1))
    return v

def update_parameters_with_momentum(W,b,tri_W,tri_b, v, beta, alpha):
    for l in range(len(nn_structure) - 1, 0, -1):
        v["dW" + str(l)] = beta * v["dW" + str(l)] + (1 - beta) * tri_W[l]
        v["db" + str(l)] = beta * v["db" + str(l)] + (1 - beta) * tri_b[l]
        # update parameters
        W[l] += -alpha * tri_W[l]
        b[l] += -alpha * tri_b[l]
    return W,b, v

def initialize_with_adam(nn_structure):
    # number of layers in the neural networks
    v = {}
    s = {}
    for l in range(1,len(nn_structure)):
        ### START CODE HERE ### (approx. 4 lines)
        v["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        v["db" + str(l)] = np.zeros((nn_structure[l],1))
        s["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        s["db" + str(l)] = np.zeros((nn_structure[l],1))
    return v, s

def update_parameters_with_adam(W,b, tri_W,tri_b, v, s, t, alpha=0.01,
                                beta1=0.9, beta2=0.999, epsilon=1e-8):
    # number of layers in the neural networks
    v_corrected = {}  # Initializing first moment estimate, python dictionary
    s_corrected = {}  # Initializing second moment estimate, python dictionary

    # Perform Adam update on all parameters
    for l in range(len(nn_structure) - 1, 0, -1):
        v["dW" + str(l)] = beta1 * v["dW" + str(l)] + (1 - beta1) * tri_W[l]
        v["db" + str(l)] = beta1 * v["db" + str(l)] + (1 - beta1) * tri_b[l]
        v_corrected["dW" + str(l)] = v["dW" + str(l)] / (1 - beta1 ** t)
        v_corrected["db" + str(l)] = v["db" + str(l)] / (1 - beta1 ** t)
        s["dW" + str(l)] = beta2 * s["dW" + str(l)] + (1 - beta2) * (tri_W[l] ** 2)
        s["db" + str(l)] = beta2 * s["db" + str(l)] + (1 - beta2) * (tri_b[l] ** 2)
        s_corrected["dW" + str(l)] = s["dW" + str(l)] / (1 - beta2 ** t)
        s_corrected["db" + str(l)] = s["db" + str(l)] / (1 - beta2 ** t)
        W[l] += - alpha * v_corrected["dW" + str(l)] / (np.sqrt(s_corrected["dW" + str(l)]) + epsilon)
        b[l] += - alpha * v_corrected["db" + str(l)] / (np.sqrt(s_corrected["db" + str(l)]) + epsilon)
    return W,b, v, s


def stratify(y, C=10):
    category = {}
    for i in range(y.shape[1]):
        for j in range(C):
            if np.argmax(y[:,i]) == j:
                category.setdefault(j,[]).append(i)# 储存类别索引
    return category

def initialize_with_ssad(nn_structure,C):
    v = {}
    s = {}
    for l in range(1, len(nn_structure)):
        for j in range(0, C):
            v["dW" + str(l) + str(j)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
            v["db" + str(l) + str(j)] = np.zeros((nn_structure[l],1))
        s["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        s["db" + str(l)] = np.zeros((nn_structure[l],1))
    return v,s

def update_parameters_with_ssag(W, b, tri_W, tri_b,v,s, alpha,C,idx):
    for l in range(len(nn_structure) - 1, 0, -1):
        # update parameters
        s["dW" + str(l)] = (s["dW" + str(l)] - v["dW" + str(l) + str(idx)])+tri_W[l]
        s["db" + str(l)] = (s["db" + str(l)] - v["db" + str(l) + str(idx)])+tri_b[l]
        v["dW" + str(l) + str(idx)] = tri_W[l]
        v["db" + str(l) + str(idx)] = tri_b[l]
        W[l] += -alpha * s["dW" + str(l)]/ C
        b[l] += -alpha * s["db" + str(l)]/ C
    return W, b,v,s

def feed_forward(x, W, b):
    h = {1: x}
    z = {}
    for l in range(1, len(W) + 1):
        if l == 1:
            node_in = x
        else:
            node_in = h[l]
        #z[l+1] = W[l].dot(node_in) + b[l] # z^(l+1) = W^(l)*h^(l) + b^(l)
        #z[l + 1] = np.dot(W[l],node_in) + b[l]
        z[l + 1] = W[l].dot(node_in) + b[l]
        if l==len(W) :
            h[l+1] = f(z[l+1])
        else:
            h[l+1] = ReluFunc(z[l+1]) # h^(l) = f(z^(l))
    return h, z
#h, z = feed_forward(x_batch, W, b)

def calculate_out_layer_delta(y, h_out, z_out):
    # delta^(nl) = -(y_i - h_i^(nl)) * f'(z_i^(nl))
    return -(y-h_out) * f_deriv(z_out)

def calculate_hidden_delta(delta_plus_1, w_l, z_l):
    # delta^(l) = (transpose(W^(l)) * delta^(l+1)) * f'(z^(l))
    return np.dot(np.transpose(w_l), delta_plus_1) * ReluPrime(z_l)
    #return np.dot(np.transpose(w_l), delta_plus_1) * f_deriv(z_l)

def train_nn_model(nn_structure, X, y,optimizer,num_epochs,bsize=1,alpha=0.0007,  C=10,lamb=0.01, beta=0.9,
                   beta1=0.9, beta2=0.999, epsilon=1e-8):
    L=len(nn_structure)
    begin_time=time()
    W, b = setup_and_init_weights(nn_structure)
    avg_cost_func = []
    t = 0
    if optimizer == "bsgd":
        pass  # no initialization required for gradient descent
    elif optimizer == "momentum":
        v = initialize_with_momentum(nn_structure)
    elif optimizer == "adam":
        v,s = initialize_with_adam(nn_structure)
    elif optimizer == "ssag":
        v,s = initialize_with_ssad(nn_structure,C)
        category = stratify(y, C=10)
        for i in range(num_epochs):
            tri_W, tri_b = init_tri_values(nn_structure)
            avg_cost = 0
            idx = np.random.randint(C)
            slice = rd.sample(list(category[idx]), bsize)  # 分类别随机抽取数据
            x_batch = X[:,slice]
            y_batch = y[:,slice]
            delta = {}
            h, z = feed_forward(x_batch, W, b)
            avg_cost = np.linalg.norm((y_batch - h[L]), axis=0).sum()
            for l in range(len(nn_structure), 0, -1):
                if l == len(nn_structure):
                    delta[l] = calculate_out_layer_delta(y_batch, h[l], z[l])
                else:
                    if l > 1:
                        delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                    tri_W[l] += np.dot(delta[l + 1], np.transpose(h[l])) * 1.0 / bsize
                    tri_b[l] += delta[l + 1] * 1.0 / bsize
            W, b, v, s = update_parameters_with_ssag(W, b, tri_W, tri_b, v, s, alpha, C, idx)
            avg_cost = 1.0 / bsize * avg_cost
            # avg_cost_func.append(avg_cost)
            if i % 100 == 0:
                print("Cost after iter %i of %i: %f" % (i, num_epochs, avg_cost))
                avg_cost_func.append(avg_cost)

    #print('Starting gradient descent for {} iterations'.format(iter_num))
    if optimizer != "ssag":
        #while cnt < iter_num:
        for i in range(num_epochs):
            #rid = np.random.randint(0, len(y) - bsize, dtype=np.int64)
            #x_batch = X[rid:rid + bsize, ...]
            #y_batch = y[rid:rid + bsize, ...]
            mini_batches= random_mini_batches(X_trains_norm, y_v_train,bsize)
            for minibatch in mini_batches:
                (x_batch, y_batch) = minibatch
                tri_W, tri_b = init_tri_values(nn_structure)
                avg_cost = 0
                delta = {}
                # 前向传播
                h, z = feed_forward(x_batch, W, b)
                # loop from nl-1 to 1 backpropagating the errors
                # 反向传播)
                avg_cost = np.linalg.norm((y_batch - h[L]), axis=0).sum()
                for l in range(len(nn_structure), 0, -1):
                    if l == len(nn_structure):
                        delta[l] = calculate_out_layer_delta(y_batch, h[l], z[l])
                    else:
                        if l > 1:
                            delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                        tri_W[l]= np.dot(delta[l + 1], np.transpose(h[l])) * 1.0 / bsize
                        tri_b[l] = np.sum(delta[l + 1],axis=1).reshape(delta[l + 1].shape[0], 1) * 1.0 / bsize
                        #grads["db" + str(L)] = np.sum(dZ, axis=1).reshape(dZ.shape[0], 1) / m
                #损失值
                avg_cost = 1.0 / bsize * avg_cost
                #avg_cost_func.append(avg_cost)
                #参数更新
                if optimizer == "bsgd":
                    W, b = update_parameters_with_bsgd(W, b, tri_W, tri_b, alpha, lamb)
                elif optimizer == "momentum":
                    W, b, v = update_parameters_with_momentum(W, b, tri_W, tri_b, v, beta, alpha)
                elif optimizer == "adam":
                    t = t + 1  # Adam counter
                    W, b, v, s = update_parameters_with_adam(W, b, tri_W, tri_b, v, s, t, alpha, beta1, beta2, epsilon)

                #print('Iteration {} of {}'.format(i, num_epochs))

            print("Cost after epochs %i  of %i: %f" % (i, num_epochs, avg_cost))
            avg_cost_func.append(avg_cost)
            # print('obj value:{} after {}th iterations at alpha:{},lamb:{}'.format(avg_cost,cnq    t,alpha,lamb))
        end_time=time()
        run_time=end_time-begin_time
        print("运行时间为：%f"  %(run_time))
    return W, b, avg_cost_func

"""
def train_nn_model(nn_structure, X, y,optimizer,num_epochs,bsize=1,alpha=0.0007,  C=10,lamb=0.01, beta=0.9,
                   beta1=0.9, beta2=0.999, epsilon=1e-8):
    W, b = setup_and_init_weights(nn_structure)
    avg_cost_func = []
    t = 0
    if optimizer == "bsgd":
        pass  # no initialization required for gradient descent
    elif optimizer == "momentum":
        v = initialize_with_momentum(nn_structure)
    elif optimizer == "adam":
        v,s = initialize_with_adam(nn_structure)
    elif optimizer == "ssag":
        v,s = initialize_with_ssad(nn_structure,C)
        category = stratify(y, C=10)
        for i in range(num_epochs):
            tri_W, tri_b = init_tri_values(nn_structure)
            avg_cost = 0
            idx = np.random.randint(C)
            slice = rd.sample(list(category[idx]), bsize)  # 分类别随机抽取数据
            x_batch = X[slice]
            y_batch = y[slice]
            for j in range(len(y_batch)):
                delta = {}
                h, z = feed_forward(x_batch[j, :], W, b)
                for l in range(len(nn_structure), 0, -1):
                    if l == len(nn_structure):
                        delta[l] = calculate_out_layer_delta(y_batch[j, :], h[l], z[l])
                        avg_cost += np.linalg.norm((y_batch[j, :] - h[l]))
                    else:
                        if l > 1:
                            delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                        tri_W[l] += np.dot(delta[l + 1][:, np.newaxis], np.transpose(h[l][:, np.newaxis])) * 1.0 / bsize
                        tri_b[l] += delta[l + 1] * 1.0 / bsize
                W, b,v,s=update_parameters_with_ssag(W, b, tri_W, tri_b, v, s, alpha, C, idx)
                avg_cost = 1.0 / bsize * avg_cost
                #avg_cost_func.append(avg_cost)
            if i % 100 == 0:
                print("Cost after iter %i of %i: %f" % (i,num_epochs, avg_cost))
                avg_cost_func.append(avg_cost)
    #print('Starting gradient descent for {} iterations'.format(iter_num))
    if optimizer != "ssag":
        #while cnt < iter_num:
        for i in range(num_epochs):
            tri_W, tri_b = init_tri_values(nn_structure)
            avg_cost = 0
            rid = np.random.randint(0, len(y) - bsize, dtype=np.int64)
            x_batch = X[rid:rid + bsize, ...]
            y_batch = y[rid:rid + bsize, ...]
            for j in range(len(y_batch)):
                delta = {}
                # 前向传播
                h, z = feed_forward(x_batch[j, :], W, b)
                # loop from nl-1 to 1 backpropagating the errors
                # 反向传播
                for l in range(len(nn_structure), 0, -1):
                    if l == len(nn_structure):
                        delta[l] = calculate_out_layer_delta(y_batch[j, :], h[l], z[l])
                        avg_cost += np.linalg.norm((y_batch[j, :] - h[l]))
                    else:
                        if l > 1:
                            delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                        tri_W[l] += np.dot(delta[l + 1][:, np.newaxis], np.transpose(h[l][:, np.newaxis])) * 1.0 / bsize
                        tri_b[l] += delta[l + 1] * 1.0 / bsize
                # 损失值
            avg_cost = 1.0 / bsize * avg_cost
            # avg_cost_func.append(avg_cost)
            # 参数更新
            if optimizer == "bsgd":
                W, b = update_parameters_with_bsgd(W, b, tri_W, tri_b, alpha, lamb)
            elif optimizer == "momentum":
                W, b, v = update_parameters_with_momentum(W, b, tri_W, tri_b, v, beta, alpha)
            elif optimizer == "adam":
                t = t + 1  # Adam counter
                W, b, v, s = update_parameters_with_adam(W, b, tri_W, tri_b, v, s, t, alpha, beta1, beta2, epsilon)

            # print('Iteration {} of {}'.format(i, num_epochs))
            if i%100==0:
                print("Cost after iter %i  of %i: %f" % (i, num_epochs, avg_cost))
                avg_cost_func.append(avg_cost)
            # print('obj value:{} after {}th iterations at alpha:{},lamb:{}'.format(avg_cost,cnt,alpha,lamb))

    return W, b, avg_cost_func


def select_parameters(X_trains_norm,y_v_train,X_tests_norm,test_labels):
    doc = docx.Document('hyperparameter.docx')
    alpha=np.arange(0.01,1,0.01).tolist()
    results=np.zeros((len(alpha),2))
    cnt=0
    for i in alpha:
        #W, b, avg_cost_func = train_nn_model(nn_structure, X_trains_norm, y_v_train, iter_num=1000, bsize=1, alpha=0.04)
        W,b,avg_cost=train_nn_model(nn_structure,X_trains_norm, y_v_train, 1000,1,i)
        #train_nn_model(nn_structure, X, y, iter_num, bsize=2, alpha=0.0007, C=10, print_cost=True)
        #train_nn_barGmst_zeromean_onesample(nn_structure, X_trains_norm, y_v_train,5000,i,j)
        yhat=predict_y(W,b,X_tests_norm,4)
        accu=accuracy_score(test_labels,yhat)*100
        results[cnt,0]=accu
        results[cnt,1]=i
        cnt+=1
    best_idx=np.argmax(results[:,0])
    doc.add_paragraph('the best hyperparameter of SSAG is alpha: {},accu:{}'.format(results[best_idx,1],results[best_idx,0]))
    print('the best hyperparameter of SSAG is alpha: {},accu:{}'.format(results[best_idx,1],results[best_idx,0]))
    doc.save('hyperparameter.docx')
    return results[best_idx,:]
"""
def predict_y(W, b, X, n_layers):
    m = X.shape[1]
    y = np.zeros((m,))
    h, z = feed_forward(X, W, b)
    for i in range(m):
        y[i] = np.argmax(h[n_layers][:,i])
    return y

#select_parameters(X_trains_norm,y_v_train,X_tests_norm,test_labels)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm,y_v_train,400,0.04,0.001)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm, y_v_train,optimizer= "bsgd",num_epochs=1000,bsize=1,alpha=0.01, lamb=0.001, beta=0.9,beta1=0.9, beta2=0.999, epsilon=1e-8)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm, y_v_train,optimizer= "momentum",num_epochs=1000,bsize=1,alpha=0.04, lamb=0.001, beta=0.9, beta1=0.9, beta2=0.999, epsilon=1e-8)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm, y_v_train,optimizer= "adam",num_epochs=100,bsize=64,alpha=0.0007, lamb=0.001, beta=0.9, beta1=0.9, beta2=0.999, epsilon=1e-8)
W, b, avg_cost_func=train_nn_model(nn_structure, X_trains_norm,y_v_train,optimizer="ssag",num_epochs=60000,bsize=1,alpha=0.04,  C=10,lamb=0.001, beta=0.9,beta1=0.9, beta2=0.999, epsilon=1e-8)

y=predict_y(W, b, X_trains_norm,4)
y_pred=predict_y(W, b, X_tests_norm,4)

print(accuracy_score(train_labels,y)*100)
print(accuracy_score(test_labels,y_pred)*100)
plt.plot(avg_cost_func)
plt.show()















#nn_structure = [784,500, 10]
#nn_structure = [784,500, 500,2000,10]
nn_structure = [784,1024,120,10]
'''
def f(x):
    return 1.0 / (1 + np.exp(-x))

'''

def sigmoid(in_x):  #  RuntimeWarning: overflow encountered in exp
    # return 1.0/(1 + np.exp(-in_x))
    # 优化方法
    if in_x >= 0:
        return 1.0/(1+np.exp(-in_x))
    else:
        return np.exp(in_x)/(1+np.exp(in_x))
'''
def f_deriv(x):
    return f(x) * (1 - f(x))
def ReluFunc(x):
    x=(np.abs(x)+x)/2.0
    return x
def ReluPrime(x):
    x[x<=0]=0
    x[x>0]=1
    return x
def LReluFunc(x):
    x[x<=0]=0.01*x
    x[x>0]=x
    return x
def LReluPrime(x):
    x[x<=0]=0.01
    x[x>0]=1
    return x

def random_mini_batches(X, Y, bsize):
    m = X.shape[0]  # number of training examples
    mini_batches = []
    # Step 1: Shuffle (X, Y)
    permutation = list(np.random.permutation(m))#np.random.permutation()随机排列序列
    shuffled_X = X[permutation,:]
    shuffled_Y = Y[permutation,:]
    # Step 2: Partition (shuffled_X, shuffled_Y). Minus the end case.
    num_complete_minibatches = math.floor(m / bsize)  # number of mini batches of size mini_batch_size in your partitionning
    for k in range(0, num_complete_minibatches):
        mini_batch_X = shuffled_X[k * bsize: (k + 1) * bsize,:]
        mini_batch_Y = shuffled_Y[k * bsize: (k + 1) * bsize,:]
        mini_batch = (mini_batch_X, mini_batch_Y)
        mini_batches.append(mini_batch)
    if m % bsize != 0:
        mini_batch_X = shuffled_X[(k + 1) * bsize:,:]
        mini_batch_Y = shuffled_Y[(k + 1) * bsize:,:]
        mini_batch = (mini_batch_X, mini_batch_Y)
        mini_batches.append(mini_batch)
    return mini_batches

def setup_and_init_weights(nn_structure):
    W = {}
    b = {}
    #np.random.seed(3)
    for l in range(1, len(nn_structure)):
        #W[l] = r.random_sample((nn_structure[l], nn_structure[l-1]))
        #b[l] = r.random_sample((nn_structure[l],))
        W[l] = np.random.randn(nn_structure[l], nn_structure[l-1])*np.sqrt(1/nn_structure[l-1])
        b[l] = np.zeros((nn_structure[l],))
    return W, b

def init_tri_values(nn_structure):
    tri_W = {}
    tri_b = {}
    for l in range(1, len(nn_structure)):
        tri_W[l] = np.zeros((nn_structure[l], nn_structure[l-1]))
        tri_b[l] = np.zeros((nn_structure[l],))
    return tri_W, tri_b

def update_parameters_with_bsgd(W, b, tri_W, tri_b, alpha, lamb):
    # Update rule for each parameter
    for l in range(len(nn_structure) - 1, 0, -1):
        W[l] += -alpha * (tri_W[l] + lamb * W[l])
        b[l] += -alpha * tri_b[l]
    return W,b

def initialize_with_momentum(nn_structure):
     # number of layers in the neural networks
    v = {}
    for l in range(1,len(nn_structure)):
        v["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        v["db" + str(l)] = np.zeros((nn_structure[l],))
    return v

def update_parameters_with_momentum(W,b,tri_W,tri_b, v, beta, alpha):
    for l in range(len(nn_structure) - 1, 0, -1):
        v["dW" + str(l)] = beta * v["dW" + str(l)] + (1 - beta) * tri_W[l]
        v["db" + str(l)] = beta * v["db" + str(l)] + (1 - beta) * tri_b[l]
        # update parameters
        W[l] += -alpha * tri_W[l]
        b[l] += -alpha * tri_b[l]
    return W,b, v

def initialize_with_adam(nn_structure):
    # number of layers in the neural networks
    v = {}
    s = {}
    for l in range(1,len(nn_structure)):
        ### START CODE HERE ### (approx. 4 lines)
        v["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        v["db" + str(l)] = np.zeros((nn_structure[l],))
        s["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        s["db" + str(l)] = np.zeros((nn_structure[l],))
    return v, s

def update_parameters_with_adam(W,b, tri_W,tri_b, v, s, t, alpha=0.01,
                                beta1=0.9, beta2=0.999, epsilon=1e-8):
    # number of layers in the neural networks
    v_corrected = {}  # Initializing first moment estimate, python dictionary
    s_corrected = {}  # Initializing second moment estimate, python dictionary

    # Perform Adam update on all parameters
    for l in range(len(nn_structure) - 1, 0, -1):
        v["dW" + str(l)] = beta1 * v["dW" + str(l)] + (1 - beta1) * tri_W[l]
        v["db" + str(l)] = beta1 * v["db" + str(l)] + (1 - beta1) * tri_b[l]
        v_corrected["dW" + str(l)] = v["dW" + str(l)] / (1 - beta1 ** t)
        v_corrected["db" + str(l)] = v["db" + str(l)] / (1 - beta1 ** t)
        s["dW" + str(l)] = beta2 * s["dW" + str(l)] + (1 - beta2) * (tri_W[l] ** 2)
        s["db" + str(l)] = beta2 * s["db" + str(l)] + (1 - beta2) * (tri_b[l] ** 2)
        s_corrected["dW" + str(l)] = s["dW" + str(l)] / (1 - beta2 ** t)
        s_corrected["db" + str(l)] = s["db" + str(l)] / (1 - beta2 ** t)
        W[l] += - alpha * v_corrected["dW" + str(l)] / (np.sqrt(s_corrected["dW" + str(l)]) + epsilon)
        b[l] += - alpha * v_corrected["db" + str(l)] / (np.sqrt(s_corrected["db" + str(l)]) + epsilon)

    return W,b, v, s


def stratify(y, C=10):
    category = {}
    for i in range(len(y)):
        for j in range(C):
            if np.argmax(y[i,:]) == j:
                category.setdefault(j,[]).append(i)# 储存类别索引
    return category
#x_batch, y_batch=stratify(X_trains_norm, y_v_train, C=10, batch=5)
#category=stratify(X_trains_norm, y_v_train, C=10, batch=1)

def initialize_with_ssad(nn_structure,C):
    v = {}
    s = {}
    for l in range(1, len(nn_structure)):
        for j in range(0, C):
            v["dW" + str(l) + str(j)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
            v["db" + str(l) + str(j)] = np.zeros((nn_structure[l],))
        s["dW" + str(l)] = np.zeros((nn_structure[l], nn_structure[l - 1]))
        s["db" + str(l)] = np.zeros((nn_structure[l],))
    return v,s

def update_parameters_with_ssag(W, b, tri_W, tri_b,v,s, alpha,C,idx):
    for l in range(len(nn_structure) - 1, 0, -1):
        # update parameters
        s["dW" + str(l)] = (s["dW" + str(l)] - v["dW" + str(l) + str(idx)])+tri_W[l]
        s["db" + str(l)] = (s["db" + str(l)] - v["db" + str(l) + str(idx)])+tri_b[l]
        v["dW" + str(l) + str(idx)] = tri_W[l]
        v["db" + str(l) + str(idx)] = tri_b[l]
        W[l] += -alpha * s["dW" + str(l)]/ C
        b[l] += -alpha * s["db" + str(l)]/ C
    return W, b,v,s

def feed_forward(x, W, b):
    h = {1: x}
    z = {}
    for l in range(1, len(W) + 1):
        if l == 1:
            node_in = x
        else:
            node_in = h[l]
        z[l+1] = W[l].dot(node_in) + b[l] # z^(l+1) = W^(l)*h^(l) + b^(l)
        if l==len(W) :
            h[l+1] = f(z[l+1])
        else:
            h[l+1] = ReluFunc(z[l+1]) # h^(l) = f(z^(l))
    return h, z

def calculate_out_layer_delta(y, h_out, z_out):
    # delta^(nl) = -(y_i - h_i^(nl)) * f'(z_i^(nl))
    return -(y-h_out) * f_deriv(z_out)

def calculate_hidden_delta(delta_plus_1, w_l, z_l):
    # delta^(l) = (transpose(W^(l)) * delta^(l+1)) * f'(z^(l))
    return np.dot(np.transpose(w_l), delta_plus_1) * ReluPrime(z_l)
    #return np.dot(np.transpose(w_l), delta_plus_1) * f_deriv(z_l)

def train_nn_model(nn_structure, X, y,optimizer,num_epochs,bsize=1,alpha=0.0007,  C=10,lamb=0.01, beta=0.9,
                   beta1=0.9, beta2=0.999, epsilon=1e-8):
    begin_time=time()
    W, b = setup_and_init_weights(nn_structure)
    avg_cost_func = []
    t = 0
    if optimizer == "bsgd":
        pass  # no initialization required for gradient descent
    elif optimizer == "momentum":
        v = initialize_with_momentum(nn_structure)
    elif optimizer == "adam":
        v,s = initialize_with_adam(nn_structure)
    elif optimizer == "ssag":
        v,s = initialize_with_ssad(nn_structure,C)
        category = stratify(y, C=10)
        for i in range(num_epochs):
            tri_W, tri_b = init_tri_values(nn_structure)
            avg_cost = 0
            idx = np.random.randint(C)
            slice = rd.sample(list(category[idx]), bsize)  # 分类别随机抽取数据
            x_batch = X[slice]
            y_batch = y[slice]
            for j in range(len(y_batch)):
                delta = {}
                h, z = feed_forward(x_batch[j, :], W, b)
                for l in range(len(nn_structure), 0, -1):
                    if l == len(nn_structure):
                        delta[l] = calculate_out_layer_delta(y_batch[j, :], h[l], z[l])
                        avg_cost += np.linalg.norm((y_batch[j, :] - h[l]))
                    else:
                        if l > 1:
                            delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                        tri_W[l] += np.dot(delta[l + 1][:, np.newaxis], np.transpose(h[l][:, np.newaxis])) * 1.0 / bsize
                        tri_b[l] += delta[l + 1] * 1.0 / bsize
                W, b,v,s=update_parameters_with_ssag(W, b, tri_W, tri_b, v, s, alpha, C, idx)
                avg_cost = 1.0 / bsize * avg_cost
                #avg_cost_func.append(avg_cost)
            if i % 100 == 0:
                print("Cost after iter %i of %i: %f" % (i,num_epochs, avg_cost))
                avg_cost_func.append(avg_cost)
    #print('Starting gradient descent for {} iterations'.format(iter_num))
    if optimizer != "ssag":
        #while cnt < iter_num:
        for i in range(num_epochs):
            #rid = np.random.randint(0, len(y) - bsize, dtype=np.int64)
            #x_batch = X[rid:rid + bsize, ...]
            #y_batch = y[rid:rid + bsize, ...]
            mini_batches= random_mini_batches(X_trains_norm, y_v_train,bsize)
            for minibatch in mini_batches:
                (x_batch, y_batch) = minibatch
                tri_W, tri_b = init_tri_values(nn_structure)
                avg_cost = 0
                for j in range(len(y_batch)):
                    delta = {}
                    #前向传播
                    h, z = feed_forward(x_batch[j, :], W, b)
                    # loop from nl-1 to 1 backpropagating the errors
                    #反向传播
                    for l in range(len(nn_structure), 0, -1):
                        if l == len(nn_structure):
                            delta[l] = calculate_out_layer_delta(y_batch[j, :], h[l], z[l])
                            avg_cost += np.linalg.norm((y_batch[j, :] - h[l]))
                        else:
                            if l > 1:
                                delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                            tri_W[l] += np.dot(delta[l + 1][:, np.newaxis],np.transpose(h[l][:, np.newaxis])) * 1.0 / bsize
                            tri_b[l] += delta[l + 1] * 1.0 / bsize
                #损失值
                avg_cost = 1.0 / bsize * avg_cost
                #avg_cost_func.append(avg_cost)
                #参数更新
                if optimizer == "bsgd":
                    W, b = update_parameters_with_bsgd(W, b, tri_W, tri_b, alpha, lamb)
                elif optimizer == "momentum":
                    W, b, v = update_parameters_with_momentum(W, b, tri_W, tri_b, v, beta, alpha)
                elif optimizer == "adam":
                    t = t + 1  # Adam counter
                    W, b, v, s = update_parameters_with_adam(W, b, tri_W, tri_b, v, s, t, alpha, beta1, beta2, epsilon)

                #print('Iteration {} of {}'.format(i, num_epochs))

            print("Cost after epochs %i  of %i: %f" % (i, num_epochs, avg_cost))
            avg_cost_func.append(avg_cost)
            # print('obj value:{} after {}th iterations at alpha:{},lamb:{}'.format(avg_cost,cnq    t,alpha,lamb))
        end_time=time()
        run_time=end_time-begin_time
        print("运行时间为：%f"  %(run_time))
    return W, b, avg_cost_func

"""
def train_nn_model(nn_structure, X, y,optimizer,num_epochs,bsize=1,alpha=0.0007,  C=10,lamb=0.01, beta=0.9,
                   beta1=0.9, beta2=0.999, epsilon=1e-8):
    W, b = setup_and_init_weights(nn_structure)
    avg_cost_func = []
    t = 0
    if optimizer == "bsgd":
        pass  # no initialization required for gradient descent
    elif optimizer == "momentum":
        v = initialize_with_momentum(nn_structure)
    elif optimizer == "adam":
        v,s = initialize_with_adam(nn_structure)
    elif optimizer == "ssag":
        v,s = initialize_with_ssad(nn_structure,C)
        category = stratify(y, C=10)
        for i in range(num_epochs):
            tri_W, tri_b = init_tri_values(nn_structure)
            avg_cost = 0
            idx = np.random.randint(C)
            slice = rd.sample(list(category[idx]), bsize)  # 分类别随机抽取数据
            x_batch = X[slice]
            y_batch = y[slice]
            for j in range(len(y_batch)):
                delta = {}
                h, z = feed_forward(x_batch[j, :], W, b)
                for l in range(len(nn_structure), 0, -1):
                    if l == len(nn_structure):
                        delta[l] = calculate_out_layer_delta(y_batch[j, :], h[l], z[l])
                        avg_cost += np.linalg.norm((y_batch[j, :] - h[l]))
                    else:
                        if l > 1:
                            delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                        tri_W[l] += np.dot(delta[l + 1][:, np.newaxis], np.transpose(h[l][:, np.newaxis])) * 1.0 / bsize
                        tri_b[l] += delta[l + 1] * 1.0 / bsize
                W, b,v,s=update_parameters_with_ssag(W, b, tri_W, tri_b, v, s, alpha, C, idx)
                avg_cost = 1.0 / bsize * avg_cost
                #avg_cost_func.append(avg_cost)
            if i % 100 == 0:
                print("Cost after iter %i of %i: %f" % (i,num_epochs, avg_cost))
                avg_cost_func.append(avg_cost)
    #print('Starting gradient descent for {} iterations'.format(iter_num))
    if optimizer != "ssag":
        #while cnt < iter_num:
        for i in range(num_epochs):
            tri_W, tri_b = init_tri_values(nn_structure)
            avg_cost = 0
            rid = np.random.randint(0, len(y) - bsize, dtype=np.int64)
            x_batch = X[rid:rid + bsize, ...]
            y_batch = y[rid:rid + bsize, ...]
            for j in range(len(y_batch)):
                delta = {}
                # 前向传播
                h, z = feed_forward(x_batch[j, :], W, b)
                # loop from nl-1 to 1 backpropagating the errors
                # 反向传播
                for l in range(len(nn_structure), 0, -1):
                    if l == len(nn_structure):
                        delta[l] = calculate_out_layer_delta(y_batch[j, :], h[l], z[l])
                        avg_cost += np.linalg.norm((y_batch[j, :] - h[l]))
                    else:
                        if l > 1:
                            delta[l] = calculate_hidden_delta(delta[l + 1], W[l], z[l])
                        tri_W[l] += np.dot(delta[l + 1][:, np.newaxis], np.transpose(h[l][:, np.newaxis])) * 1.0 / bsize
                        tri_b[l] += delta[l + 1] * 1.0 / bsize
                # 损失值
            avg_cost = 1.0 / bsize * avg_cost
            # avg_cost_func.append(avg_cost)
            # 参数更新
            if optimizer == "bsgd":
                W, b = update_parameters_with_bsgd(W, b, tri_W, tri_b, alpha, lamb)
            elif optimizer == "momentum":
                W, b, v = update_parameters_with_momentum(W, b, tri_W, tri_b, v, beta, alpha)
            elif optimizer == "adam":
                t = t + 1  # Adam counter
                W, b, v, s = update_parameters_with_adam(W, b, tri_W, tri_b, v, s, t, alpha, beta1, beta2, epsilon)

            # print('Iteration {} of {}'.format(i, num_epochs))
            if i%100==0:
                print("Cost after iter %i  of %i: %f" % (i, num_epochs, avg_cost))
                avg_cost_func.append(avg_cost)
            # print('obj value:{} after {}th iterations at alpha:{},lamb:{}'.format(avg_cost,cnt,alpha,lamb))

    return W, b, avg_cost_func


def select_parameters(X_trains_norm,y_v_train,X_tests_norm,test_labels):
    doc = docx.Document('hyperparameter.docx')
    alpha=np.arange(0.01,1,0.01).tolist()
    results=np.zeros((len(alpha),2))
    cnt=0
    for i in alpha:
        #W, b, avg_cost_func = train_nn_model(nn_structure, X_trains_norm, y_v_train, iter_num=1000, bsize=1, alpha=0.04)
        W,b,avg_cost=train_nn_model(nn_structure,X_trains_norm, y_v_train, 1000,1,i)
        #train_nn_model(nn_structure, X, y, iter_num, bsize=2, alpha=0.0007, C=10, print_cost=True)
        #train_nn_barGmst_zeromean_onesample(nn_structure, X_trains_norm, y_v_train,5000,i,j)
        yhat=predict_y(W,b,X_tests_norm,4)
        accu=accuracy_score(test_labels,yhat)*100
        results[cnt,0]=accu
        results[cnt,1]=i
        cnt+=1
    best_idx=np.argmax(results[:,0])
    doc.add_paragraph('the best hyperparameter of SSAG is alpha: {},accu:{}'.format(results[best_idx,1],results[best_idx,0]))
    print('the best hyperparameter of SSAG is alpha: {},accu:{}'.format(results[best_idx,1],results[best_idx,0]))
    doc.save('hyperparameter.docx')
    return results[best_idx,:]
"""
def predict_y(W, b, X, n_layers):
    m = X.shape[0]
    y = np.zeros((m,))
    for i in range(m):
        h, z = feed_forward(X[i, :], W, b)
        y[i] = np.argmax(h[n_layers])
    return y
#select_parameters(X_trains_norm,y_v_train,X_tests_norm,test_labels)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm,y_v_train,400,0.04,0.001)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm, y_v_train,optimizer= "bsgd",num_epochs=1000,bsize=1,alpha=0.01, lamb=0.001, beta=0.9,beta1=0.9, beta2=0.999, epsilon=1e-8)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm, y_v_train,optimizer= "momentum",num_epochs=1000,bsize=1,alpha=0.04, lamb=0.001, beta=0.9, beta1=0.9, beta2=0.999, epsilon=1e-8)
#W, b, avg_cost_func=train_nn_model(nn_structure,X_trains_norm, y_v_train,optimizer= "adam",num_epochs=10,bsize=64,alpha=0.0007, lamb=0.001, beta=0.9, beta1=0.9, beta2=0.999, epsilon=1e-8)
W, b, avg_cost_func=train_nn_model(nn_structure, X_trains_norm,y_v_train,optimizer="ssag",num_epochs=1000,bsize=1,alpha=0.04,  C=10,lamb=0.001, beta=0.9,beta1=0.9, beta2=0.999, epsilon=1e-8)

y=predict_y(W, b, X_trains_norm,4)
y_pred=predict_y(W, b, X_tests_norm,4)

print(accuracy_score(train_labels,y)*100)
print(accuracy_score(test_labels,y_pred)*100)
plt.plot(avg_cost_func)
plt.show()
'''
